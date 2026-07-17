import argparse
from pathlib import Path
from docx import Document  # type: ignore
#from docx.text.run import Run

def paragraph_to_markdown(paragraph, ImageFolder: Path, ImageIndex: list[int]) -> str:
    if paragraph.style.name.startswith("Heading"):
        try:
            Level = int(paragraph.style.name.split()[-1])
        except Exception:
            Level = 1

        Prefix = "#" * Level + " "
    else:
        Prefix = ""

    Parts = []

    CurrentText = ""
    CurrentBold = None
    CurrentItalic = None
    
    def flush():
        nonlocal CurrentText, CurrentBold, CurrentItalic

        if not CurrentText:
            return

        if CurrentBold and CurrentItalic:
            Parts.append(f"***{CurrentText}***")
        elif CurrentBold:
            Parts.append(f"**{CurrentText}**")
        elif CurrentItalic:
            Parts.append(f"*{CurrentText}*")
        else:
            Parts.append(CurrentText)

        CurrentText = ""


    for Run in paragraph.runs:  # noqa: F402
        Drawing = Run._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")

        if Drawing is not None:
            Blip = Drawing.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")

            if Blip is not None:
                Rid = Blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")

                if Rid:
                    ImagePart = Run.part.related_parts[Rid]

                    Ext = ImagePart.content_type.split("/")[-1]

                    ImageName = f"IMG{ImageIndex[0]:03d}.{Ext}"
                    ImagePath = ImageFolder / ImageName

                    ImagePath.write_bytes(ImagePart.blob)

                    flush()
                    Parts.append(f"[IMG{ImageIndex[0]:03d}]")

                    ImageIndex[0] += 1

                    continue
        if not Run.text:
            continue

        Bold = bool(Run.bold)
        Italic = bool(Run.italic)

        if CurrentBold is None:
            CurrentBold = Bold
            CurrentItalic = Italic

        if Bold == CurrentBold and Italic == CurrentItalic:
            CurrentText += Run.text
        else:
            flush()
            CurrentBold = Bold
            CurrentItalic = Italic
            CurrentText = Run.text

    flush()

    return Prefix + "".join(Parts)


def convert_file(input_file: Path, output_file: Path):
    print(f"[START] {input_file}")

    document = Document(str(input_file))
    ImageFolder = output_file.with_suffix("")
    ImageFolder.mkdir(parents=True, exist_ok=True)

    ImageIndex = [1]

    total = len(document.paragraphs)

    with output_file.open("w", encoding="utf-8", newline="\n") as fout:
        for index, paragraph in enumerate(document.paragraphs, 1):
            fout.write(
    paragraph_to_markdown(
        paragraph,
        ImageFolder,
        ImageIndex
    )
)
            fout.write("\n")

            if index % 500 == 0 or index == total:
                print(f"  {index}/{total} paragraphs\r", end="")

    print()
    print(f"[DONE] {output_file}\n")


def batch_convert(input_folder: Path, output_folder: Path):
    docx_files = sorted(input_folder.glob("*.docx"))

    if not docx_files:
        print("Không tìm thấy file docx")
        return

    output_folder.mkdir(parents=True, exist_ok=True)

    total = len(docx_files)

    print(f"Batch start: {total} files\n")

    for index, docx_file in enumerate(docx_files, 1):
        print(f"=== File {index}/{total} ===")

        output_file = output_folder / (docx_file.stem + ".txt")
        convert_file(docx_file, output_file)

    print("Batch completed")


def main():
    parser = argparse.ArgumentParser()

    parser.add_argument(
        "mode",
        choices=["single", "batch"],
        help="single = xử lý 1 file, batch = xử lý cả thư mục",
    )

    parser.add_argument(
        "input",
        help="file docx hoặc folder",
    )

    parser.add_argument(
        "-o",
        "--output",
        help="file hoặc folder output",
    )

    args = parser.parse_args()

    input_path = Path(args.input)

    if args.mode == "single":
        output_path = Path(args.output or input_path.with_suffix(".txt"))
        convert_file(input_path, output_path)

    else:
        output_folder = Path(args.output or "output_txt")
        batch_convert(input_path, output_folder)


if __name__ == "__main__":
    main()