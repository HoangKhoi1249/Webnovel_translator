import argparse
from pathlib import Path
from docx import Document


def convert_file(InputFile: Path, OutputFile: Path):
    print(f"[START] {InputFile}")

    DocumentObj = Document()

    with open(InputFile, "r", encoding="utf-8") as f:
        Lines = f.readlines()

    TotalLines = len(Lines)

    for Index, Line in enumerate(Lines, 1):
        Line = Line.rstrip()

        if not Line:
            DocumentObj.add_paragraph("")
            continue

        if Line.startswith("#"):
            Level = len(Line) - len(Line.lstrip("#"))
            Text = Line[Level:].strip()

            DocumentObj.add_heading(Text, level=min(Level, 9))
        else:
            DocumentObj.add_paragraph(Line)

        if Index % 500 == 0 or Index == TotalLines:
            print(f"  {Index}/{TotalLines} lines")

    DocumentObj.save(str(OutputFile))

    print(f"[DONE] {OutputFile}\n")


def batch_convert(InputFolder: Path, OutputFolder: Path):
    TxtFiles = sorted(InputFolder.glob("*.txt"))
    TotalFiles = len(TxtFiles)

    if TotalFiles == 0:
        print("Không tìm thấy file txt")
        return

    OutputFolder.mkdir(parents=True, exist_ok=True)

    print(f"Batch start: {TotalFiles} files\n")

    for Index, TxtFile in enumerate(TxtFiles, 1):
        print(f"=== File {Index}/{TotalFiles} ===")

        OutputFile = OutputFolder / (TxtFile.stem + ".docx")
        convert_file(TxtFile, OutputFile)

    print("Batch completed")


def main():
    Parser = argparse.ArgumentParser()

    Parser.add_argument(
        "mode",
        choices=["single", "batch"],
        help="single = xử lý 1 file, batch = xử lý cả thư mục"
    )

    Parser.add_argument(
        "input",
        help="file txt hoặc folder"
    )

    Parser.add_argument(
        "-o",
        "--output",
        help="file hoặc folder output"
    )

    Args = Parser.parse_args()

    InputPath = Path(Args.input)

    if Args.mode == "single":
        OutputPath = Path(Args.output or InputPath.with_suffix(".docx"))
        convert_file(InputPath, OutputPath)

    elif Args.mode == "batch":
        OutputFolder = Path(Args.output or "output_docx")
        batch_convert(InputPath, OutputFolder)


if __name__ == "__main__":
    main()